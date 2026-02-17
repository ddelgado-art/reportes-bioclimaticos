"""
================================================================================
  GENERADOR DE REPORTES BIOCLIMÁTICOS — Streamlit v3 (app_final_v3.py)
  Autor base: Auditoría QA comparativa EPW_file_v9.ipynb
  Correcciones v3 aplicadas sobre app_auditada_v2.py:

  [FIX-V3-01] 'sector' vs 'sector_viento': CRÍTICO — agregada clave 'sector'
              al context (DOCX usa {{ sector }}, no {{ sector_viento }})
  [FIX-V3-02] Race condition tempfiles: nombres únicos por sesión con uuid
  [FIX-V3-03] os.remove() movido a bloque finally para cleanup garantizado
  [FIX-V3-04] encoding='utf-8' agregado a DB_estaciones_meteo_procesado.csv
  [FIX-V3-05] Dead import mpatches eliminado
  [FIX-V3-06] except Exception: continue → muestra warning al usuario en tablas
  [FIX-V3-07] EPW malformado capturado con st.error() + st.stop()
  [FIX-V3-08] Timeout ArcGIS capturado con aviso al usuario
  [FIX-V3-09] .streamlit/config.toml comentado para referencia
================================================================================

CORRECCIONES PREVIAS MANTENIDAS (de app_auditada_v2.py):
  [FIX-A1] Normalización proyecto_normalizado eliminaba comas → corregido
  [FIX-A2] Validación de campo proyecto vacío → corregido
  [FIX-B1] Municipio no extraído del notebook → alineado
  [FIX-C1] texto_epw_completo incluye pais_epw → corregido
  [FIX-C2] Confort adaptativo usa epw_copy → alineado
  [FIX-C3] Gráfica confort OOP vs state machine → alineado visualmente
  [FIX-D1] Tablas DOCX sobre doc.docx post-render → corregido
  [FIX-D2] Firma insertar_subdocumento_memoria unificada → corregido
  [FIX-E1] Plantilla buscada como 'Plantilla_Reporte.docx' → corregido
"""
try:
    import pkg_resources
except ImportError:
    import setuptools
    import pkg_resources


import streamlit as st
import pandas as pd
import numpy as np
import geopandas as gpd
from shapely.geometry import Point
import zipfile
import io
import os
import uuid       # [FIX-V3-02] Para nombres únicos de archivos temporales
import requests
import math
import unicodedata
from xml.etree import ElementTree as ET
from pvlib import iotools
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from docx import Document
from copy import deepcopy
from pythermalcomfort.models import adaptive_ashrae
from pythermalcomfort.utilities import running_mean_outdoor_temperature
import matplotlib
matplotlib.use('Agg')          # ✅ ANTES de cualquier import de pyplot
import matplotlib.pyplot as plt
import matplotlib.colors as mcolors
import matplotlib.dates as mdates
# [FIX-V3-05] mpatches eliminado — no se usa en ninguna parte de la app
import seaborn as sns
from windrose import WindroseAxes
from PIL import Image, ImageDraw

# ==========================================
# 0. CONFIGURACIÓN Y ESTILOS GLOBALES
# ==========================================
st.set_page_config(page_title="Reportes Bioclimáticos", layout="wide")

# ✅ Idéntico al notebook Cell 43 — whitegrid + edgecolor 0.15 + gridcolor 0.75
sns.set_theme(style="whitegrid", rc={
    "axes.edgecolor": "0.15",
    "grid.color": "0.75",
    "grid.linestyle": "-"
})

# ==========================================
# 1. FUNCIONES CORE
# ==========================================

def corregir_coordenadas(val):
    """Corrige coordenadas con múltiples puntos decimales (ej: -74.125.883 → -74.125883).
    Idéntico al notebook Cell 35."""
    if isinstance(val, str):
        partes = val.split('.')
        if len(partes) > 2:
            return float(partes[0] + '.' + ''.join(partes[1:]))
        return float(val)
    return float(val)


def calcular_distancia(lat1, lon1, lat2, lon2):
    """Distancia Haversine en km. Idéntico al notebook Cell 35."""
    rad = math.pi / 180
    dlat = (lat2 - lat1) * rad
    dlon = (lon2 - lon1) * rad
    a = math.sin(dlat / 2) ** 2 + math.cos(lat1 * rad) * math.cos(lat2 * rad) * math.sin(dlon / 2) ** 2
    c = 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))
    return 6371.0 * c


def grados_a_direccion(grados):
    """Convierte grados brújula a texto cardinal. Idéntico al notebook Cell 37."""
    val = int((grados / 45) + 0.5)
    direcciones = ["norte", "noreste", "este", "sureste", "sur", "suroeste", "oeste", "noroeste"]
    return direcciones[(val % 8)]


@st.cache_data
def extraer_coordenadas_kmz(kmz_content: bytes):
    """Extrae latitud y longitud del primer punto encontrado en un archivo KMZ.
    Idéntico al notebook Cell 21 — namespace kml con fallback sin namespace.
    @st.cache_data: argumento bytes es hasheable → correcto."""
    try:
        with zipfile.ZipFile(io.BytesIO(kmz_content)) as kmz:
            kml_file = [name for name in kmz.namelist() if name.endswith('.kml')][0]
            kml_content = kmz.read(kml_file)
        root = ET.fromstring(kml_content)
        ns = {'kml': 'http://www.opengis.net/kml/2.2'}
        coords = root.find('.//kml:coordinates', ns)
        if coords is None:
            coords = root.find('.//coordinates')
        if coords is not None and coords.text:
            coord_text = coords.text.strip().split()[0]
            lon_val, lat_val = coord_text.split(',')[:2]
            return float(lat_val), float(lon_val)
    except Exception:
        pass
    return None, None


def obtener_mapa_arcgis_memoria(lat_val, lon_val, zoom):
    """Descarga mapa satelital ArcGIS y dibuja diana concéntrica en memoria.
    [FIX-V3-08] Timeout y errores HTTP notificados al usuario.
    ✅ Misma URL, parámetros, tamaño y DPI del notebook Cell 45.
    ✅ Diana con 6 círculos (radios 30,25,20,15,10,5; colores white/red/white/red/white/red)."""
    width, height, dpi = 1200, 900, 300
    scale_factor = 156543.04 / (2 ** zoom)
    delta_lat = (height / 2 * scale_factor) / 111320
    delta_lon = (width / 2 * scale_factor) / (111320 * 0.5)
    bbox = f"{lon_val - delta_lon},{lat_val - delta_lat},{lon_val + delta_lon},{lat_val + delta_lat}"
    url = (
        f"https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/export"
        f"?bbox={bbox}&bboxSR=4326&size={width},{height}&imageSR=4326"
        f"&format=png24&transparent=false&dpi={dpi}&f=image"
    )
    try:
        response = requests.get(url, timeout=25)
        if response.status_code == 200:
            img = Image.open(io.BytesIO(response.content))
            if img.mode != 'RGB':
                img = img.convert('RGB')
            draw = ImageDraw.Draw(img)
            center_x, center_y = width // 2, height // 2
            # Diana: 6 círculos concéntricos rojo/blanco — idéntica al notebook Cell 45
            radios = [30, 25, 20, 15, 10, 5]
            colores = ['white', 'red', 'white', 'red', 'white', 'red']
            for radio, color in zip(radios, colores):
                draw.ellipse(
                    [center_x - radio, center_y - radio, center_x + radio, center_y + radio],
                    fill=color
                )
            img_byte_arr = io.BytesIO()
            img.save(img_byte_arr, format='PNG', quality=100, optimize=False)
            img_byte_arr.seek(0)
            return img_byte_arr
        else:
            st.warning(f"⚠️ ArcGIS respondió con código {response.status_code}. Mapa no disponible.")
    except requests.Timeout:
        # [FIX-V3-08] Timeout explícito con mensaje al usuario
        st.warning("⏱️ Timeout al descargar mapa ArcGIS (>25s). El informe se generará sin imagen satelital.")
    except Exception as e:
        st.warning(f"⚠️ Error descargando mapa satelital: {e}")
    return None


def insertar_subdocumento_memoria(doc_base, texto_marcador, ruta_anexo):
    """Inserta contenido de un DOCX anexo en la posición del marcador de texto.
    [FIX-D2] Firma unificada: (doc_base_objeto, texto_marcador, ruta_anexo).
    Equivalente al notebook Cell 30 operando en memoria."""
    if not os.path.exists(ruta_anexo):
        return doc_base
    doc_anexo = Document(ruta_anexo)
    parrafo_marcador = None
    for p in doc_base.paragraphs:
        if texto_marcador in p.text:
            parrafo_marcador = p
            break
    if parrafo_marcador is not None:
        p_base_xml = parrafo_marcador._element
        for elemento in doc_anexo.element.body:
            if not elemento.tag.endswith('sectPr'):
                p_base_xml.addprevious(elemento)
        p_base_xml.getparent().remove(p_base_xml)
    return doc_base


def cargar_texto_condicional(ruta):
    """Lee texto de un archivo si existe. Retorna '' si no existe (nunca lanza excepción).
    Idéntico al notebook Cell 31 (cargar_texto)."""
    if os.path.exists(ruta):
        with open(ruta, 'r', encoding='utf-8') as f:
            return f.read()
    return ""


# ==========================================
# 2. UI STREAMLIT — SIDEBAR
# ==========================================
st.title("📊 Generador de Reportes Bioclimáticos")
st.sidebar.header("Inputs Generales")

# Variables generales del proyecto — equivalentes al notebook Cell 9
input_proyecto = st.sidebar.text_input("Proyecto", "SIA Boyacá2")
input_ubicacion = st.sidebar.text_input("Ubicación", "Sogamoso, Boyacá")
input_cliente = st.sidebar.text_input("Cliente", "Consorcio SIA Boyacá")

# Cargar tipos de edificios — equivalente al notebook Cell 17
try:
    df_edificios = pd.read_csv('TiposEdificios.csv', sep=';', encoding='latin-1')
    opciones_edificios_display = [
        f"{row['Tipo']} ({row['ID']})" for _, row in df_edificios.iterrows()
    ]
    opciones_edificios_ids = df_edificios['ID'].tolist()
    opciones_map = dict(zip(opciones_edificios_display, opciones_edificios_ids))
except FileNotFoundError:
    st.error("❌ Falta TiposEdificios.csv en el repositorio.")
    st.stop()

seleccionados_display = st.sidebar.multiselect(
    "Tipos de Edificio (ASHRAE):",
    options=opciones_edificios_display
)
ids_seleccionados = [opciones_map[d] for d in seleccionados_display]

# Cargar alcances — equivalente al notebook Cell 12
try:
    with open('alcances_tipo.txt', 'r', encoding='utf-8') as archivo:
        opciones_alcance = [linea.strip() for linea in archivo.readlines() if linea.strip()]
except FileNotFoundError:
    opciones_alcance = ["Opción por defecto"]

alcance_seleccionado = st.sidebar.selectbox("Alcance:", opciones_alcance)

# Carga de archivos — equivalente al notebook Cell 21
epw_file = st.sidebar.file_uploader("1. Cargar EPW", type=['epw'])
kmz_file = st.sidebar.file_uploader("2. Cargar KMZ/KML", type=['kmz', 'kml'])

# ==========================================
# 3. EJECUCIÓN DEL MOTOR
# ==========================================
if st.sidebar.button("🚀 Generar Análisis y Ensamblar Informe"):

    # ✅ Validación de entradas completa [FIX-A2]
    errores = []
    if not input_proyecto.strip():
        errores.append("El nombre del proyecto no puede estar vacío.")
    if not epw_file:
        errores.append("Debes cargar un archivo EPW.")
    if not ids_seleccionados:
        errores.append("Debes seleccionar al menos un tipo de edificio.")
    if errores:
        for e in errores:
            st.warning(f"⚠️ {e}")
        st.stop()

    with st.spinner("⚙️ Procesando análisis bioclimático..."):

        # ----------------------------------------------------------
        # [FIX-A1] Normalización de proyecto_normalizado
        # Elimina tildes + espacios → _ + comas
        # ----------------------------------------------------------
        proyecto_normalizado = (
            unicodedata.normalize('NFKD', input_proyecto)
            .encode('ASCII', 'ignore')
            .decode('utf-8')
            .replace(' ', '_')
            .replace(',', '')
        )

        # ----------------------------------------------------------
        # Texto de tipos de edificio — idéntico al notebook Cell 19
        # ----------------------------------------------------------
        datos_seleccionados = df_edificios[df_edificios['ID'].isin(ids_seleccionados)]
        lista_tipos = datos_seleccionados['Tipo'].str.lower().tolist()
        cantidad = len(lista_tipos)
        if cantidad == 0:
            texto_tipos = ""
        elif cantidad == 1:
            texto_tipos = lista_tipos[0]
        elif cantidad == 2:
            texto_tipos = f"{lista_tipos[0]} y {lista_tipos[1]}"
        else:
            texto_tipos = ", ".join(lista_tipos[:-1]) + f" y {lista_tipos[-1]}"
        texto_ids = ", ".join(ids_seleccionados)

        # ----------------------------------------------------------
        # 3.1 Leer EPW — [FIX-V3-07] validación de archivo malformado
        # ----------------------------------------------------------
        try:
            epw_data, epw_metadata = iotools.read_epw(epw_file, coerce_year=2024)
        except Exception as epw_err:
            st.error(f"❌ El archivo EPW no es válido o está dañado: {epw_err}")
            st.stop()

        lat = epw_metadata['latitude']
        lon = epw_metadata['longitude']

        if kmz_file is not None:
            lat_k, lon_k = extraer_coordenadas_kmz(kmz_file.getvalue())
            if lat_k is not None and lon_k is not None:
                lat, lon = lat_k, lon_k
                st.info(f"🎯 Usando coordenadas del KMZ: {lat:.4f}, {lon:.4f}")
            else:
                st.info(f"🎯 KMZ sin coordenadas válidas. Usando coordenadas del EPW: {lat:.4f}, {lon:.4f}")
        else:
            st.info(f"🎯 Usando coordenadas del EPW: {lat:.4f}, {lon:.4f}")

        # ----------------------------------------------------------
        # 3.2 Geoprocesamiento — idéntico al notebook Cell 27
        # ----------------------------------------------------------
        punto_proyecto = Point(lon, lat)
        es_bogota = False
        departamento_proyecto = "No Identificado"
        municipio_proyecto = "No Identificado"
        codigo_dane_proyecto = "00"
        norma_1 = ""
        norma_2 = ""

        try:
            mapa_bogota = gpd.read_file('perimetrourbano.json')
            if mapa_bogota.crs != "EPSG:4326":
                mapa_bogota = mapa_bogota.to_crs("EPSG:4326")
            if mapa_bogota.contains(punto_proyecto).any():
                es_bogota = True
                departamento_proyecto = "Bogotá D.C."
                municipio_proyecto = "Bogotá D.C."
                norma_1 = "Decreto 582 de 2023 – Ecourbanismo y Construcción Sostenible, Bogotá."
                norma_2 = "Decreto 099 de 2024 - Normas de construcción y habitabilidad aplicables a la vivienda urbana y rural en Bogotá."
        except Exception as geo_err:
            st.warning(f"⚠️ Error al cargar mapa Bogotá: {geo_err}")

        if not es_bogota:
            try:
                gdf = gpd.read_file('colombia.geo.json')
                if gdf.crs is None:
                    gdf.set_crs(epsg=4326, inplace=True)
                coincidencia = gdf[gdf.contains(punto_proyecto)]
                if not coincidencia.empty:
                    fila = coincidencia.iloc[0]
                    departamento_proyecto = fila.get('NOMBRE_DPT', 'Desconocido')
                    municipio_proyecto = fila.get('NOMBRE_MPI', fila.get('NOMBRE_DPT', 'Desconocido'))
                    codigo_dane_proyecto = str(fila.get('DPTO', '00'))
                else:
                    gdf['distancia'] = gdf.distance(punto_proyecto)
                    mejor = gdf.sort_values('distancia').iloc[0]
                    departamento_proyecto = mejor.get('NOMBRE_DPT', 'Desconocido')
                    municipio_proyecto = mejor.get('NOMBRE_MPI', mejor.get('NOMBRE_DPT', 'Desconocido'))
                    codigo_dane_proyecto = str(mejor.get('DPTO', '00'))
            except Exception as geo_err2:
                st.warning(f"⚠️ Error al cargar mapa nacional: {geo_err2}")

        # ----------------------------------------------------------
        # Estación meteorológica más cercana — idéntico al notebook Cell 35
        # ----------------------------------------------------------
        est_codigo = "N/A"
        est_nombre = "N/A"
        est_municipio = "N/A"
        est_departamento_var = "N/A"
        est_altitud = "N/A"
        est_longitud = "N/A"
        est_latitud = "N/A"
        est_distancia = "N/A"

        try:
            # [FIX-V3-04] Agregado encoding='utf-8' para nombres con caracteres latinos
            df_estaciones = pd.read_csv(
                'DB_estaciones_meteo_procesado.csv', sep=';', encoding='utf-8'
            )
            df_estaciones['LATITUD_DECIMAL'] = df_estaciones['LATITUD_DECIMAL'].apply(corregir_coordenadas)
            df_estaciones['LONGITUD_DECIMAL'] = df_estaciones['LONGITUD_DECIMAL'].apply(corregir_coordenadas)
            df_estaciones['distancia_km'] = df_estaciones.apply(
                lambda row: calcular_distancia(lat, lon, row['LATITUD_DECIMAL'], row['LONGITUD_DECIMAL']),
                axis=1
            )
            estacion_cercana = df_estaciones.loc[df_estaciones['distancia_km'].idxmin()]
            est_codigo = str(estacion_cercana['CODIGO'])
            est_nombre = estacion_cercana['NOMBRE']
            est_municipio = estacion_cercana['MUNICIPIO']
            est_departamento_var = estacion_cercana['DEPARTAMENTO']
            est_altitud = f"{estacion_cercana['ELEVACIÓN']} m.s.n.m."
            est_longitud = estacion_cercana['LONGITUD']
            est_latitud = estacion_cercana['LATITUD']
            est_distancia = f"{estacion_cercana['distancia_km']:.2f} km"
        except Exception as est_err:
            st.warning(f"⚠️ Error buscando estación meteorológica: {est_err}")

        # ----------------------------------------------------------
        # 3.3 Control Normativo y de Escenarios — notebook Cells 10 + 31
        # ----------------------------------------------------------
        es_edificio_medico = 'MED' in ids_seleccionados
        año_referencia = epw_data.index.year.unique()[0]
        es_futuro = año_referencia > 2025

        ctx_cond = {
            'epw_future_1': cargar_texto_condicional('t_EPWfuture_text1_1.txt') if es_futuro else "",
            'epw_future_2': (
                cargar_texto_condicional('t_EPWfuture_text2_1.txt') if es_futuro
                else cargar_texto_condicional('t_EPWfuture_text2_2.txt')
            ),
            'natvent_1': (
                cargar_texto_condicional('t_natvent_text1_MED.txt') if es_edificio_medico
                else cargar_texto_condicional('t_natvent_text1_GEN.txt')
            ),
            'natvent_2': "",
            'bogota_1': cargar_texto_condicional('t_bogota_text1.txt') if es_bogota else "",
            'bogota_2': cargar_texto_condicional('t_bogota_text2.txt') if es_bogota else "",
            'bogota_3': "",
        }

        # ----------------------------------------------------------
        # 3.4 DataFrame climático — idéntico al notebook Cell 33
        # ----------------------------------------------------------
        df_clima = epw_data[['temp_air', 'relative_humidity', 'wind_direction', 'wind_speed']].copy()
        df_clima.reset_index(inplace=True)
        df_clima.rename(columns={
            df_clima.columns[0]: 'Tiempo',
            'temp_air': 'Temp_Bulbo_Seco',
            'relative_humidity': 'Humedad_Relativa',
            'wind_direction': 'Dir_Viento',
            'wind_speed': 'Vel_Viento'
        }, inplace=True)

        # ----------------------------------------------------------
        # Variables estadísticas climáticas — idéntico al notebook Cell 37
        # ----------------------------------------------------------
        site_name = epw_metadata.get('city', 'N/A')
        pais_epw = epw_metadata.get('country', 'No especificado')   # [FIX-C1]
        altura_msnm = epw_metadata.get('altitude', 'N/A')
        lat_epw = epw_metadata.get('latitude', 0)
        lon_epw = epw_metadata.get('longitude', 0)
        alt_epw = epw_metadata.get('altitude', 0)

        t_min = df_clima['Temp_Bulbo_Seco'].min()
        t_prom = df_clima['Temp_Bulbo_Seco'].mean()
        t_max = df_clima['Temp_Bulbo_Seco'].max()
        t_mas_2_val = t_prom + 2
        t_mas_4_val = t_prom + 4
        t_menos_2_val = t_prom - 2
        t_menos_4_val = t_prom - 4

        h_min = df_clima['Humedad_Relativa'].min()
        h_prom = df_clima['Humedad_Relativa'].mean()
        h_max = df_clima['Humedad_Relativa'].max()

        v_min = df_clima['Vel_Viento'].min()
        v_prom = df_clima['Vel_Viento'].mean()
        v_max = df_clima['Vel_Viento'].max()

        dir_predominante = df_clima['Dir_Viento'].mode()[0]
        d_inicio_val = (dir_predominante - 45) % 360
        d_fin_val = (dir_predominante + 45) % 360
        sector_texto = (
            f"{grados_a_direccion(d_inicio_val)}-"
            f"{grados_a_direccion(dir_predominante)}-"
            f"{grados_a_direccion(d_fin_val)}"
        )

        if año_referencia > 2025:
            tipo_clima = "futuro"
            complemento_texto = (
                f"Para garantizar la resiliencia del diseño ante el cambio climático, "
                f"se generó una proyección a futuro bajo el escenario del IPCC con un RCP 4.5 "
                f"proyectado al año {año_referencia}."
            )
        else:
            tipo_clima = "contemporaneo"
            complemento_texto = "Este archivo representa un año meteorológico típico (TMY) basado en datos históricos contemporáneos."

        texto_epw_completo = (
            f"Los datos climáticos utilizados provienen del archivo EPW correspondiente a "
            f"{site_name}, {pais_epw}. La estación meteorológica se ubica en las coordenadas "
            f"{lat_epw:.4f}°N, {lon_epw:.4f}°W, a una altitud de {alt_epw:.0f} metros sobre el nivel del mar."
        )

        # ----------------------------------------------------------
        # Confort térmico: percentiles + ASHRAE — idéntico al notebook Cell 39
        # ----------------------------------------------------------
        t_90_inf_val = df_clima['Temp_Bulbo_Seco'].quantile(0.05)
        t_90_sup_val = df_clima['Temp_Bulbo_Seco'].quantile(0.95)
        t_80_inf_val = df_clima['Temp_Bulbo_Seco'].quantile(0.10)
        t_80_sup_val = df_clima['Temp_Bulbo_Seco'].quantile(0.90)

        # [FIX-C2] usa epw_data.copy() igual al notebook Cell 39
        epw_copy = epw_data.copy()
        epw_copy['mes'] = epw_copy.index.month
        temp_mensual = epw_copy.groupby('mes')['temp_air'].mean()
        t_rm = temp_mensual.mean()

        resultado_confort = adaptive_ashrae(tdb=t_prom, tr=t_prom, t_running_mean=t_rm, v=0.1)
        t_90_inf_ashrae = resultado_confort['tmp_cmf_90_low']
        t_90_sup_ashrae = resultado_confort['tmp_cmf_90_up']
        t_80_inf_ashrae = resultado_confort['tmp_cmf_80_low']
        t_80_sup_ashrae = resultado_confort['tmp_cmf_80_up']

        # ----------------------------------------------------------
        # DataFrames ASHRAE — idéntico al notebook Cell 41
        # ----------------------------------------------------------
        columnas_requeridas = ['Categoria_Es', 'RP_Ls', 'Ra_Ls'] + ids_seleccionados
        df_ashrae = pd.read_csv(
            'ASH621_Tabla 6-1_types.csv', sep=';', encoding='utf-8',
            skiprows=[0], usecols=columnas_requeridas, na_values=['missing value']
        )
        mascara_filtro = pd.Series(False, index=df_ashrae.index)
        for id_tipo in ids_seleccionados:
            col_str = df_ashrae[id_tipo].astype(str).str.strip().str.lower()
            mascara_filtro |= (col_str != 'nan') & (col_str != '') & (col_str != 'none')
        df_tasas_filtradas = df_ashrae[mascara_filtro].copy().reset_index(drop=True)
        df_tasas_word = df_tasas_filtradas[['Categoria_Es', 'RP_Ls', 'Ra_Ls']].fillna("")

        columnas_requeridas_62 = ['Categoria_Es', 'Exh_Ls_u', 'Exh_Ls_a'] + ids_seleccionados
        variaciones_nulos = ['missing value', 'Missing Value', 'Missing value', ' ', '-']
        df_ashrae_62 = pd.read_csv(
            'ASH621_Tabla 6-2_types.csv', sep=';', encoding='utf-8',
            skiprows=[0], usecols=columnas_requeridas_62, na_values=variaciones_nulos
        )
        mascara_filtro_62 = pd.Series(False, index=df_ashrae_62.index)
        for id_tipo in ids_seleccionados:
            col_str = df_ashrae_62[id_tipo].astype(str).str.strip().str.lower()
            mascara_filtro_62 |= (col_str != 'nan') & (col_str != '') & (col_str != 'none')
        df_extr = df_ashrae_62[mascara_filtro_62].copy().reset_index(drop=True)
        df_extr['Exh_Ls_u'] = pd.to_numeric(df_extr['Exh_Ls_u'], errors='coerce')
        df_extr['Exh_Ls_a'] = pd.to_numeric(df_extr['Exh_Ls_a'], errors='coerce')
        df_extr = df_extr.fillna({'Exh_Ls_u': 0.0, 'Exh_Ls_a': 0.0})
        df_extr_word = df_extr[['Categoria_Es', 'Exh_Ls_u', 'Exh_Ls_a']].fillna("")

        # ----------------------------------------------------------
        # 3.5 Mapas satelitales — idéntico al notebook Cell 45
        # ----------------------------------------------------------
        mapa_gen_stream = obtener_mapa_arcgis_memoria(lat, lon, 14)
        mapa_pun_stream = obtener_mapa_arcgis_memoria(lat, lon, 20)

        # ----------------------------------------------------------
        # 3.6 Gráfica de Temperatura — idéntico al notebook Cell 47
        # ----------------------------------------------------------
        fig_temp_stream = None
        try:
            fig, ax = plt.subplots(figsize=(12, 4))
            ax.fill_between(
                df_clima['Tiempo'],
                df_clima['Temp_Bulbo_Seco'].rolling(window=24, min_periods=1).min(),
                df_clima['Temp_Bulbo_Seco'].rolling(window=24, min_periods=1).max(),
                color='#ff4d4d', alpha=0.2, label='Rango Diario (Mín-Máx)'
            )
            ax.plot(
                df_clima['Tiempo'],
                df_clima['Temp_Bulbo_Seco'].rolling(window=168, min_periods=1).mean(),
                color='#cc0000', linewidth=2, label='Promedio Semanal'
            )
            meses_es = ['Ene', 'Feb', 'Mar', 'Abr', 'May', 'Jun', 'Jul', 'Ago', 'Sep', 'Oct', 'Nov', 'Dic']
            ax.xaxis.set_major_locator(mdates.MonthLocator())
            ax.set_xticks(pd.date_range(df_clima['Tiempo'].min(), df_clima['Tiempo'].max(), freq='MS'))
            ax.set_xticklabels(meses_es)
            ax.set_title('Temperatura de Bulbo Seco (°C) - Comportamiento Anual',
                         loc='left', fontweight='bold', fontsize=11, pad=20)
            ax.legend(loc='lower right', bbox_to_anchor=(1.0, 1.02), ncol=2, frameon=False, fontsize=9)
            ax.set_ylabel('Temp (°C)')
            ax.set_xlabel('')
            ax.set_xlim(df_clima['Tiempo'].min(), df_clima['Tiempo'].max())
            plt.tight_layout()
            fig_temp_stream = io.BytesIO()
            fig.savefig(fig_temp_stream, format='PNG', dpi=300, bbox_inches="tight")
            fig_temp_stream.seek(0)
            plt.close(fig)
        except Exception as e:
            st.warning(f"⚠️ Error generando gráfica de temperatura: {e}")

        # ----------------------------------------------------------
        # 3.7 Gráfica de Humedad — idéntico al notebook Cell 49
        # ----------------------------------------------------------
        fig_hum_stream = None
        try:
            fig, ax = plt.subplots(figsize=(12, 4))
            ax.fill_between(
                df_clima['Tiempo'],
                df_clima['Humedad_Relativa'].rolling(window=24, min_periods=1).min(),
                df_clima['Humedad_Relativa'].rolling(window=24, min_periods=1).max(),
                color='#0066cc', alpha=0.15, label='Rango Diario (Mín-Máx)'
            )
            ax.plot(
                df_clima['Tiempo'],
                df_clima['Humedad_Relativa'].rolling(window=168, min_periods=1).mean(),
                color='#003366', linewidth=2, label='Promedio Semanal'
            )
            meses_es = ['Ene', 'Feb', 'Mar', 'Abr', 'May', 'Jun', 'Jul', 'Ago', 'Sep', 'Oct', 'Nov', 'Dic']
            ax.xaxis.set_major_locator(mdates.MonthLocator())
            ax.set_xticks(pd.date_range(df_clima['Tiempo'].min(), df_clima['Tiempo'].max(), freq='MS'))
            ax.set_xticklabels(meses_es)
            ax.set_title('Humedad Relativa (%) - Análisis Anual',
                         loc='left', fontweight='bold', fontsize=11, pad=20)
            ax.legend(loc='lower right', bbox_to_anchor=(1.0, 1.02), ncol=2, frameon=False, fontsize=9)
            ax.set_ylabel('Humedad (%)')
            ax.set_xlabel('')
            ax.set_xlim(df_clima['Tiempo'].min(), df_clima['Tiempo'].max())
            ax.grid(True, axis='both', zorder=0)
            plt.tight_layout()
            fig_hum_stream = io.BytesIO()
            fig.savefig(fig_hum_stream, format='PNG', dpi=300, bbox_inches="tight")
            fig_hum_stream.seek(0)
            plt.close(fig)
        except Exception as e:
            st.warning(f"⚠️ Error generando gráfica de humedad: {e}")

        # ----------------------------------------------------------
        # 3.8 Rosa de Vientos — idéntico al notebook Cell 51
        # ----------------------------------------------------------
        fig_viento_stream = None
        try:
            colores_viento = ['#A5E7FA', '#4BD3F7', '#2D9CDB', '#2F52E0',
                              '#7B2FBE', '#C82396', '#E02F2F', '#F2994A']
            cmap_custom = mcolors.ListedColormap(colores_viento)
            bins_velocidad = [0, 0.5, 1.5, 3.3, 5.5, 7.9, 10.7, 13.8]
            fig = plt.figure(figsize=(8, 8))
            ax_wind = WindroseAxes.from_ax(fig=fig)
            ax_wind.bar(
                df_clima['Dir_Viento'], df_clima['Vel_Viento'],
                normed=True, opening=0.85, edgecolor='white',
                cmap=cmap_custom, bins=bins_velocidad, nsector=16
            )
            ax_wind.set_legend(
                title="Velocidad (m/s)", loc='lower right',
                bbox_to_anchor=(1.2, 0), decimal_places=1
            )
            plt.title('Rosa de los Vientos Anual', fontsize=14, fontweight='bold', pad=20)
            plt.tight_layout()
            fig_viento_stream = io.BytesIO()
            fig.savefig(fig_viento_stream, format='PNG', dpi=300, bbox_inches="tight")
            fig_viento_stream.seek(0)
            plt.close(fig)
        except Exception as e:
            st.warning(f"⚠️ Error generando rosa de vientos: {e}")

        # ----------------------------------------------------------
        # 3.9 Gráfica de Confort Adaptativo ASHRAE 55
        # [FIX-C3] OOP (fig/ax) alineado visualmente con notebook Cell 53
        # ----------------------------------------------------------
        fig_confort_stream = None
        try:
            df_diario = epw_data['temp_air'].resample('D').mean()
            t_pma = running_mean_outdoor_temperature(df_diario.values, alpha=0.8)
            if isinstance(t_pma, (float, np.float64, np.float32)):
                t_pma = [t_pma]

            fig, ax = plt.subplots(figsize=(10, 7))
            x_axis = np.linspace(10, 33.5, 100)
            y_neutral = 0.31 * x_axis + 17.8
            ax.fill_between(x_axis, y_neutral - 3.5, y_neutral + 3.5,
                            color='#ffcc80', alpha=0.4, label='80% Aceptabilidad')
            ax.fill_between(x_axis, y_neutral - 2.5, y_neutral + 2.5,
                            color='#a5d6a7', alpha=0.5, label='90% Aceptabilidad')
            ax.plot(x_axis, y_neutral, color='#666666', linestyle='--', label='Neutralidad')
            ax.set_title('Análisis de Confort Adaptativo', fontsize=12, fontweight='bold', loc='left')
            ax.set_xlabel('Temp. Media Exterior Predominante (°C)', fontsize=10)
            ax.set_ylabel('Temp. Operativa Interior (°C)', fontsize=10)
            ax.legend(loc='upper right', frameon=True, fontsize=9)
            ax.grid(True, linestyle=':', alpha=0.6, color='gray', linewidth=0.5)
            plt.tight_layout()
            fig_confort_stream = io.BytesIO()
            fig.savefig(fig_confort_stream, format='PNG', dpi=300, bbox_inches="tight")
            fig_confort_stream.seek(0)
            plt.close(fig)
        except Exception as e:
            st.warning(f"⚠️ Error generando gráfica de confort: {e}")

        # ----------------------------------------------------------
        # 3.10 Previsualización en Streamlit
        # ----------------------------------------------------------
        col1, col2 = st.columns(2)
        if mapa_gen_stream:
            mapa_gen_stream.seek(0)
            col1.image(mapa_gen_stream.read(), caption="Mapa General (zoom 14)", use_container_width=True)
            mapa_gen_stream.seek(0)
        if mapa_pun_stream:
            mapa_pun_stream.seek(0)
            col2.image(mapa_pun_stream.read(), caption="Mapa Puntual (zoom 20)", use_container_width=True)
            mapa_pun_stream.seek(0)

        if fig_temp_stream:
            fig_temp_stream.seek(0)
            st.image(fig_temp_stream.read(), caption="Temperatura Anual", use_container_width=True)
            fig_temp_stream.seek(0)
        if fig_hum_stream:
            fig_hum_stream.seek(0)
            st.image(fig_hum_stream.read(), caption="Humedad Relativa", use_container_width=True)
            fig_hum_stream.seek(0)

        col3, col4 = st.columns(2)
        if fig_viento_stream:
            fig_viento_stream.seek(0)
            col3.image(fig_viento_stream.read(), caption="Rosa de Vientos", use_container_width=True)
            fig_viento_stream.seek(0)
        if fig_confort_stream:
            fig_confort_stream.seek(0)
            col4.image(fig_confort_stream.read(), caption="Confort Adaptativo ASHRAE 55", use_container_width=True)
            fig_confort_stream.seek(0)

        # ----------------------------------------------------------
        # 3.11 Ensamblaje DOCX
        # [FIX-E1] Nombre correcto: 'Plantilla_Reporte.docx'
        # ----------------------------------------------------------
        nombre_plantilla = 'Plantilla_Reporte.docx'
        if not os.path.exists(nombre_plantilla):
            nombre_plantilla = 'plantilla_base.docx'   # fallback legacy

        if not os.path.exists(nombre_plantilla):
            st.error(f"❌ No se encontró la plantilla Word ({nombre_plantilla}). Verifica el repositorio.")
            st.stop()

        # [FIX-V3-02] Nombres únicos por sesión — evita race conditions con usuarios simultáneos
        session_id = str(uuid.uuid4())[:8]
        img_paths = {
            'img_general':      (f'mapa_general_{session_id}.png',         mapa_gen_stream,    100),
            'img_puntual':      (f'mapa_puntual_{session_id}.png',         mapa_pun_stream,    100),
            'img_temp_anual':   (f'grafica_temperatura_{session_id}.png',  fig_temp_stream,    150),
            'img_hum_anual':    (f'grafica_humedad_{session_id}.png',      fig_hum_stream,     150),
            'img_rosa_vientos': (f'grafica_rosa_vientos_{session_id}.png', fig_viento_stream,  120),
            'img_confort':      (f'grafica_confort_{session_id}.png',      fig_confort_stream, 140),
        }

        # Escribir imágenes a disco temporalmente para DocxTemplate
        temp_imgs = {}
        try:
            doc = DocxTemplate(nombre_plantilla)

            for key, (fname, stream, ancho_mm) in img_paths.items():
                if stream is not None:
                    stream.seek(0)
                    with open(fname, 'wb') as f:
                        f.write(stream.read())
                    temp_imgs[key] = fname

            # Contexto DOCX — equivalente al notebook Cell 59 + Cell 60
            # [FIX-V3-01] CRÍTICO: clave 'sector' agregada (DOCX usa {{ sector }})
            context = {
                # Proyecto
                'proyecto': input_proyecto,
                'ubicacion': input_ubicacion,
                'cliente': input_cliente,
                'nombre_documento': f"{proyecto_normalizado}_InformeBioclimatico",
                'tipos_edificio': texto_tipos,
                'tipo_alcance': alcance_seleccionado,
                # Ubicación geográfica
                'departamento': departamento_proyecto,
                'municipio': municipio_proyecto,
                'latitud': f"{lat:.3f}",
                'longitud': f"{lon:.3f}",
                'alt': str(altura_msnm),
                # Normativas Bogotá (en context para uso futuro en plantilla)
                'norma_1': norma_1,
                'norma_2': norma_2,
                # Temperatura
                't_min': f"{t_min:.1f}",
                't_prom': f"{t_prom:.1f}",
                't_max': f"{t_max:.1f}",
                't_mas_2': f"{t_mas_2_val:.1f}",
                't_mas_4': f"{t_mas_4_val:.1f}",
                't_menos_2': f"{t_menos_2_val:.1f}",
                't_menos_4': f"{t_menos_4_val:.1f}",
                # Humedad
                'h_min': f"{h_min:.1f}",
                'h_prom': f"{h_prom:.1f}",
                'h_max': f"{h_max:.1f}",
                # Viento
                'd_inicio': int(d_inicio_val),
                'd_fin': int(d_fin_val),
                'v_min': f"{v_min:.1f}",
                'v_prom': f"{v_prom:.1f}",
                'v_max': f"{v_max:.1f}",
                # [FIX-V3-01] CORRECCIÓN CRÍTICA: 'sector' requerido por {{ sector }} del DOCX
                'sector': sector_texto,
                'sector_viento': sector_texto,     # alias legacy por si se actualiza la plantilla
                # Confort estadístico
                't_90_inf': f"{t_90_inf_val:.1f}",
                't_90_sup': f"{t_90_sup_val:.1f}",
                't_80_inf': f"{t_80_inf_val:.1f}",
                't_80_sup': f"{t_80_sup_val:.1f}",
                # Confort ASHRAE 55 adaptativo (disponibles si la plantilla los incorpora)
                't_90_inf_ashrae': f"{t_90_inf_ashrae:.1f}",
                't_90_sup_ashrae': f"{t_90_sup_ashrae:.1f}",
                't_80_inf_ashrae': f"{t_80_inf_ashrae:.1f}",
                't_80_sup_ashrae': f"{t_80_sup_ashrae:.1f}",
                # Texto EPW [FIX-C1]
                'info_proyeccion': texto_epw_completo,
                # Estación meteorológica más cercana
                'est_codigo': est_codigo,
                'est_nombre': est_nombre,
                'est_municipio': est_municipio,
                'est_depto': est_departamento_var,
                'est_altitud': est_altitud,
                'est_long': str(est_longitud),
                'est_lat': str(est_latitud),
                'est_dist': est_distancia,
                # Textos condicionales — equivalente al .update() del notebook Cell 60
                **ctx_cond,
            }

            # Inyectar InlineImage al contexto
            for key, (fname, stream, ancho_mm) in img_paths.items():
                if key in temp_imgs and os.path.exists(temp_imgs[key]):
                    context[key] = InlineImage(doc, temp_imgs[key], width=Mm(ancho_mm))

            # Renderizar plantilla
            doc.render(context)

            # [FIX-D1] Llenado de tablas ASHRAE sobre doc.docx (post-render, pre-save)
            doc_obj = doc.docx

            for table in doc_obj.tables:
                try:
                    if len(table.rows) < 2:
                        continue
                    header_text = " ".join(cell.text for cell in table.rows[0].cells)

                    # Tabla 6-1: Renovación de aire
                    if 'Categoría' in header_text and 'RP' in header_text and 'Ra' in header_text:
                        template_row = table.rows[1]
                        while len(table.rows) > 1:
                            table._element.remove(table.rows[-1]._element)
                        for _, row_data in df_tasas_word.iterrows():
                            new_row_element = deepcopy(template_row._element)
                            table._element.append(new_row_element)
                            new_row = table.rows[-1]
                            for cell_idx, valor in enumerate(
                                [row_data['Categoria_Es'], row_data['RP_Ls'], row_data['Ra_Ls']]
                            ):
                                cell = new_row.cells[cell_idx]
                                if cell.paragraphs and cell.paragraphs[0].runs:
                                    cell.paragraphs[0].runs[0].text = str(valor)
                                else:
                                    cell.text = str(valor)

                    # Tabla 6-2: Extracción de aire
                    elif 'Categoría' in header_text and 'Exh' in header_text:
                        template_row = table.rows[1]
                        while len(table.rows) > 1:
                            table._element.remove(table.rows[-1]._element)
                        for _, row_data in df_extr_word.iterrows():
                            new_row_element = deepcopy(template_row._element)
                            table._element.append(new_row_element)
                            new_row = table.rows[-1]
                            for cell_idx, valor in enumerate(
                                [row_data['Categoria_Es'], row_data['Exh_Ls_u'], row_data['Exh_Ls_a']]
                            ):
                                cell = new_row.cells[cell_idx]
                                if cell.paragraphs and cell.paragraphs[0].runs:
                                    cell.paragraphs[0].runs[0].text = str(valor)
                                else:
                                    cell.text = str(valor)

                except Exception as tbl_err:
                    # [FIX-V3-06] Muestra advertencia en lugar de silenciar el error
                    st.warning(f"⚠️ Error llenando tabla DOCX: {tbl_err}")
                    continue

            # Guardar al stream intermedio
            doc_stream = io.BytesIO()
            doc.save(doc_stream)
            doc_stream.seek(0)

            # Acoplar Anexos en memoria [FIX-D2]
            doc_final = Document(doc_stream)
            if es_edificio_medico:
                doc_final = insertar_subdocumento_memoria(
                    doc_final, '<<<ANEXO_MED>>>', 't_natvent_text2_MED.docx'
                )
            if es_bogota:
                doc_final = insertar_subdocumento_memoria(
                    doc_final, '<<<ANEXO_BOGOTA>>>', 't_bogota_text3.docx'
                )

            # Limpiar marcadores residuales
            marcadores_a_limpiar = ['<<<ANEXO_MED>>>', '<<<ANEXO_BOGOTA>>>']
            for p in list(doc_final.paragraphs):
                if any(m in p.text for m in marcadores_a_limpiar):
                    p._element.getparent().remove(p._element)

            # Guardar el documento final en memoria
            doc_out_stream = io.BytesIO()
            doc_final.save(doc_out_stream)

            st.session_state['docx_data'] = doc_out_stream.getvalue()
            st.session_state['docx_nombre'] = f"{proyecto_normalizado}_InformeBioclimatico.docx"

        finally:
            # [FIX-V3-03] Limpieza GARANTIZADA de archivos temporales en bloque finally
            for fname in temp_imgs.values():
                try:
                    if os.path.exists(fname):
                        os.remove(fname)
                except Exception:
                    pass

    st.success(
        "✅ Análisis bioclimático, normativas ASHRAE, gráficas y documento ensamblados correctamente."
    )

# ==========================================
# 4. BOTÓN DE DESCARGA PERSISTENTE
# ==========================================
if 'docx_data' in st.session_state:
    st.download_button(
        label="📥 Descargar Informe Bioclimático (DOCX)",
        data=st.session_state['docx_data'],
        file_name=st.session_state.get('docx_nombre', 'InformeBioclimatico.docx'),
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
